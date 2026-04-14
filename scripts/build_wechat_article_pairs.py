#!/usr/bin/env python3
"""
Fill `sop 测试母版-模板.docx`-style meeting template with two public WeChat article excerpts;
write SOP (`rewrite_paragraph`) vs anti-pattern (`paragraph.text`) pairs.

Usage (from repo root):
  python scripts/build_wechat_article_pairs.py \\
    --template "/path/to/sop 测试母版-模板.docx" \\
    --out-dir demo/out-wechat
"""
from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document

BODY_PARA_START = 7
BODY_CHUNKS = 6


def rewrite_paragraph(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def bad_assign(paragraph, new_text: str) -> None:
    paragraph.text = new_text


def split_body(text: str, parts: int = BODY_CHUNKS, soft_limit: int = 780) -> list[str]:
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in paras:
        if len(buf) + len(p) + 2 <= soft_limit:
            buf = f"{buf}\n\n{p}" if buf else p
        else:
            if buf:
                chunks.append(buf)
                buf = p
            else:
                chunks.append(p[: soft_limit + 120])
                buf = ""
        if len(chunks) >= parts:
            break
    if buf and len(chunks) < parts:
        chunks.append(buf)
    while len(chunks) < parts:
        chunks.append("")
    return chunks[:parts]


def apply_meta_and_body(doc: Document, sop: bool, meta: dict, body_chunks: list[str]) -> None:
    assign = rewrite_paragraph if sop else bad_assign
    assign(doc.paragraphs[0], meta["headline"])
    assign(doc.paragraphs[1], meta["file_line"])
    assign(doc.paragraphs[3], meta["doc_title"])
    assign(doc.paragraphs[4], meta["source_line"])
    assign(doc.paragraphs[6], "【正文摘录】")
    for i, chunk in enumerate(body_chunks):
        idx = BODY_PARA_START + i
        if idx >= len(doc.paragraphs):
            break
        if sop:
            rewrite_paragraph(doc.paragraphs[idx], chunk)
        else:
            bad_assign(doc.paragraphs[idx], chunk)


def write_pair(template: Path, out_dir: Path, slug: str, meta: dict, body: str) -> tuple[Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    chunks = split_body(body)
    bad_p = out_dir / f"{slug}-bad-paragraph-text.docx"
    sop_p = out_dir / f"{slug}-sop-rewrite-paragraph.docx"
    shutil.copy(template, bad_p)
    shutil.copy(template, sop_p)
    d1 = Document(str(bad_p))
    apply_meta_and_body(d1, False, meta, chunks)
    d1.save(str(bad_p))
    d2 = Document(str(sop_p))
    apply_meta_and_body(d2, True, meta, chunks)
    d2.save(str(sop_p))
    return bad_p, sop_p


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", type=Path, required=True)
    ap.add_argument("--out-dir", type=Path, default=Path("demo/out-wechat"))
    args = ap.parse_args()
    if not args.template.is_file():
        raise SystemExit(f"template not found: {args.template}")

    repo = Path(__file__).resolve().parents[1]
    src = repo / "demo/article-sources"

    articles: list[tuple[str, dict[str, str], str]] = [
        (
            "01-liberal-arts",
            {
                "headline": "公众号正文入库演示（同一母版）",
                "file_line": "【演示】杂文节选 · 文科与 AI",
                "doc_title": "文档：AI 时代，鼓吹文科生有用的是纯扯淡",
                "source_line": "来源：https://mp.weixin.qq.com/s/zoG31kifchxAizD5dSbpAA",
            },
            (src / "01-liberal-arts.txt").read_text(encoding="utf-8"),
        ),
        (
            "02-rag-local",
            {
                "headline": "公众号正文入库演示（同一母版）",
                "file_line": "【演示】长文节选 · 本地 RAG",
                "doc_title": "文档：基于本地的个人RAG数据库实践：效果不尽如人意",
                "source_line": "来源：https://mp.weixin.qq.com/s/uCtNiviw11VtvhzV2__zJQ",
            },
            (src / "02-rag-local.txt").read_text(encoding="utf-8"),
        ),
    ]

    for slug, meta, body in articles:
        write_pair(args.template, args.out_dir, slug, meta, body)
        print("Wrote pair:", slug)


if __name__ == "__main__":
    main()
