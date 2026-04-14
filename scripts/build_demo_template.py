#!/usr/bin/env python3
"""Build a small formatted .docx to feed compare_sop_vs_paragraph_text.py (demo only)."""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def set_ea_font(run, west: str, east: str) -> None:
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", type=Path, default=Path("demo/demo-template.docx"))
    args = ap.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 0 — title (compare script rewrites this paragraph entirely)
    p0 = doc.add_paragraph()
    r = p0.add_run("慢行街区与公共空间（版式演示母版）")
    r.bold = True
    r.font.size = Pt(22)
    set_ea_font(r, "Arial", "黑体")

    # 1 — spacer (unchanged by compare script)
    doc.add_paragraph("")

    # 2 — greeting: normal + bold runs
    p2 = doc.add_paragraph()
    r1 = p2.add_run("各位读者：")
    set_ea_font(r1, "Times New Roman", "宋体")
    r1.font.size = Pt(12)
    r2 = p2.add_run("（本句应为加粗，用于对照）")
    set_ea_font(r2, "Times New Roman", "宋体")
    r2.font.size = Pt(12)
    r2.bold = True

    # 3–7 — body with two runs (will show merge to one run vs keep first run only)
    for i in range(3, 8):
        p = doc.add_paragraph()
        a = p.add_run(f"【待替换段落 {i}】正文前段宋体 ")
        set_ea_font(a, "Times New Roman", "宋体")
        a.font.size = Pt(12)
        b = p.add_run("后段加粗强调")
        set_ea_font(b, "Times New Roman", "宋体")
        b.font.size = Pt(12)
        b.bold = True

    doc.save(str(args.out))
    print("Wrote", args.out.resolve())


if __name__ == "__main__":
    main()
