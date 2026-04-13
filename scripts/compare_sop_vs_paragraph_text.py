#!/usr/bin/env python3
"""
Generate two .docx from the SAME template for side-by-side comparison:
  - sop: rewrite_paragraph (runs[0] + clear other runs) — preserves rPr on first run
  - bad: paragraph.text = ... — often destroys run-level formatting

Usage:
  python compare_sop_vs_paragraph_text.py --template /path/to/template.docx --out-dir ./out

Do NOT run inside another git repo if you want a clean publish folder; this script only writes files.
"""
from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document


def rewrite_paragraph(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def bad_assign(paragraph, new_text: str) -> None:
    paragraph.text = new_text


# Neutral demo text (no org-specific content)
BLOCKS = {
    0: "慢行街区与公共空间：一篇短文（对照示例）",
    2: "各位读者：",
    3: (
        "城市更新里，最容易被看见的是路面颜色与护栏样式；更难的是路口转弯半径、路缘抬升与公交站视线。"
        "这些细节决定一条路是“能走”，还是“愿意反复走”。"
    ),
    4: (
        "治理层面，慢行友好往往不是靠“禁止机动车”，而是把路权规则写清楚：哪些时段共享、哪些路段限速、哪些节点必须留白。"
        "规则可预期，冲突就会下降。"
    ),
    5: (
        "对市民而言，收益常常体现在时间结构上：短途不必绑定小汽车；对沿街小店而言，慢行流量更可能带来停留。"
        "因此试点街区常把遮阳躲雨与停留设施写进同一套清单。"
    ),
    6: (
        "落地时总会遇到现实约束：地下管线、消防通道、夜间照明、无障碍坡道。"
        "好的方案通常允许迭代：先消除最危险的不一致，再逐步补齐舒适与美观。"
    ),
    7: (
        "如果把慢行系统理解为城市操作系统的一次补丁升级，评价标准就不只是“有没有”，而是“好不好用、好不好维护、好不好解释”。"
        "谢谢阅读。"
    ),
}


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", type=Path, required=True, help="Path to an existing formatted .docx (母版)")
    ap.add_argument("--out-dir", type=Path, default=Path("out"), help="Output directory")
    args = ap.parse_args()

    tpl: Path = args.template
    if not tpl.is_file():
        raise SystemExit(f"template not found: {tpl}")

    out_dir: Path = args.out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    sop_path = out_dir / "compare-sop-rewrite-paragraph.docx"
    bad_path = out_dir / "compare-bad-paragraph-text.docx"

    shutil.copy(tpl, sop_path)
    doc = Document(str(sop_path))
    for idx, text in BLOCKS.items():
        if idx < len(doc.paragraphs):
            rewrite_paragraph(doc.paragraphs[idx], text)
    doc.save(str(sop_path))

    shutil.copy(tpl, bad_path)
    doc2 = Document(str(bad_path))
    for idx, text in BLOCKS.items():
        if idx < len(doc2.paragraphs):
            bad_assign(doc2.paragraphs[idx], text)
    doc2.save(str(bad_path))

    print("Wrote:")
    print(" ", sop_path)
    print(" ", bad_path)


if __name__ == "__main__":
    main()
